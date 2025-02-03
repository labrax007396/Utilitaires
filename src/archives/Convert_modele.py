
from onnx.tools.net_drawer import GetPydotGraph, GetOpNodeProducer
import onnxruntime as rt
from onnxmltools import convert_lightgbm
from skl2onnx import to_onnx, update_registered_converter
from skl2onnx.common.shape_calculator import calculate_linear_regressor_output_shapes  # noqa
from onnxmltools import __version__ as oml_version
from onnxmltools.convert.lightgbm.operator_converters.LightGbm import convert_lightgbm
from onnxruntime.capi.onnxruntime_pybind11_state import Fail as OrtFail
from skl2onnx import convert_sklearn, update_registered_converter
from skl2onnx.common.shape_calculator import calculate_linear_classifier_output_shapes  # noqa
from onnxmltools.convert.lightgbm.operator_converters.LightGbm import convert_lightgbm  # noqa
import onnxmltools.convert.common.data_types
from skl2onnx.common.data_types import FloatTensorType, StringTensorType
from skl2onnx.common.data_types import Int64TensorType
from onnx.onnx_pb import StringStringEntryProto
from lightgbm import LGBMRegressor





# INFORMATIONS GENERALES

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def convert_to_docx(modelreport_json, dico_model, dir_case):
        
    import json
    from docx import Document
    from docx.shared import Mm

    json_rep = json.loads(modelreport_json)
    # create document
    doc = Document()

    doc.add_heading("Informations générales",level=1)

    # add grid table
    table_info = doc.add_table(rows=6, cols=2, style="Table Grid")

    des_row = table_info.rows[0].cells
    des_row[0].text = "Site"
    des_row[1].text = json_rep['site']

    var_row = table_info.rows[1].cells
    var_row[0].text = "Variable modélisée"
    var_row[1].text = dico_model['description']

    var_per = table_info.rows[2].cells
    var_per[0].text = "Période"
    var_per[1].text = "Du " + json_rep['dataframe_info']['start_date'] + " Au " + json_rep['dataframe_info']['end_date']


    var_nbp = table_info.rows[3].cells
    var_nbp[0].text = "Nombre de points"
    var_nbp[1].text = str(json_rep['dataframe_info']['cleaning_info']['line_count_after'])


    var_nbf = table_info.rows[4].cells
    var_nbf[0].text = "Nombre de facteurs retenus"
    var_nbf[1].text = str(len([t['nom'] for t in dico_model['facteurs'].values() if t['used']]))


    var_freq = table_info.rows[5].cells
    var_freq[0].text = "Fréquence"
    var_freq[1].text = dico_model['freq']


    # FACTEURS

    doc.add_heading("Facteurs",level=1)

    # add grid table



    table_fac = doc.add_table(rows=len(dico_model['facteurs'])+1, cols=4, style="Table Grid")

    heading_row = table_fac.rows[0].cells

    # add headings
    heading_row[0].text = "Alias"
    heading_row[1].text = "Description"
    heading_row[2].text = "Unité"
    heading_row[3].text = "Facteur retenu"


    num_row = 0
    for f in dico_model['facteurs'].values():
        if f['used']:
            num_row = num_row + 1
            var_fact = table_fac.rows[num_row].cells
            var_fact[0].text = f['nom']
            var_fact[1].text = f['description']
            var_fact[2].text = f['unit']
            var_fact[3].text = 'Oui'

    for f in dico_model['facteurs'].values():
        if not(f['used']):
            num_row = num_row + 1
            var_fact = table_fac.rows[num_row].cells
            var_fact[0].text = f['nom']
            var_fact[1].text = f['description']
            var_fact[2].text = f['unit']
            var_fact[3].text = 'Non'            

    make_rows_bold(table_fac.rows[0])

    
    doc.add_heading('Distribution facteurs numériques', 2)
    doc.add_picture(dir_case + 'distrib fact num.png', width=Mm(150))
    
    doc.add_heading('Répartition facteurs discrets', 2)
    doc.add_picture(dir_case + 'distrib fact cat.png', width=Mm(150))

    doc.add_heading('Coefficients de Corrélation', 2)
    doc.add_picture(dir_case + 'coefficient correlation.png', width=Mm(150))
    # MODELE

    doc.add_heading("Modèle",level=1)

    # add grid table
    table_mod = doc.add_table(rows=4, cols=2, style="Table Grid")

    mod_row = table_mod.rows[0].cells
    mod_row[0].text = "Type de modèle"
    mod_row[1].text = json_rep['model_info']['model_type']

    r2_row = table_mod.rows[1].cells
    r2_row[0].text = "Coefficient détermination (R2)"
    r2_row[1].text = str(round(json_rep['model_info']['r2_train'],2))

    mape_row = table_mod.rows[2].cells
    mape_row[0].text = "Erreur moyenne relative (%)"
    mape_row[1].text = str(round(json_rep['model_info']['mape_train'],2))

    form_row = table_mod.rows[3].cells
    form_row[0].text = "Formule"
    form_row[1].text = json_rep['model_info']['formula']

    doc.add_heading('Modèle .vs. Mesure', 2)
    doc.add_picture(dir_case + 'model vs mesure.png', width=Mm(150))

    # figure


    doc.add_heading('Graphes', 1)
   


    #p = doc.add_paragraph()
    #r = p.add_run()
    #r.add_picture(dir_case + 'coefficient correlation.png')

    return doc



















def skl2onnx_convert_lightgbm(scope, operator, container):
    options = scope.get_options(operator.raw_operator)
    if 'split' in options:
        if pv.Version(oml_version) < pv.Version('1.9.2'):
            warnings.warn(
                "Option split was released in version 1.9.2 but %s is "
                "installed. It will be ignored." % oml_version)
        operator.split = options['split']
    else:
        operator.split = None
    convert_lightgbm(scope, operator, container)

def convert_to_onnx(X, model, modelreport_json):

    update_registered_converter(
        LGBMRegressor, 'LightGbmLGBMRegressor',
        calculate_linear_regressor_output_shapes,
        skl2onnx_convert_lightgbm,
        options={'split': None})
    
    inputs = []
    for k, v in zip(X.columns, X.dtypes):
        if v == 'int64':
            t = Int64TensorType([None, 1])
        elif v == 'float64':
            t = FloatTensorType([None, 1])
        else:
            t = StringTensorType([None, 1])
        inputs.append((k, t))

    output = [('target',FloatTensorType([None, 1]))]

    model_onnx = to_onnx(model, initial_types=inputs,final_types=output,
                        target_opset={'': 13, 'ai.onnx.ml': 2})

    model_onnx.metadata_props.append(StringStringEntryProto(key="ReportModel", value = modelreport_json))

    return model_onnx