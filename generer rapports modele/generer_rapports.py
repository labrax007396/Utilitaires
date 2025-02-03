
def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True



def word_docx(dico_infos:dict):

    from docx import Document
    from docx.shared import Mm
    from docx.shared import Inches, Pt,Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ### entête


    doc = Document()
    header = doc.sections[0].header
    htable=header.add_table(1, 2, Inches(6))
    
    htab_cells=htable.rows[0].cells
    ht0=htab_cells[1].add_paragraph()

    kh=ht0.add_run()
    kh.add_picture('logouw.png', width=Inches(2))
    ht0.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ht1=htab_cells[0].add_paragraph('Rapport modélisation')
    ht1.alignment = WD_ALIGN_PARAGRAPH.LEFT


    make_rows_bold(htable.rows[0])
 
    ### infos générales

    json_rep   = dico_infos['rapport']
    dico_model = dico_infos['config']
    doc.add_heading("Informations générales",level=1)


    table_info = doc.add_table(rows=3, cols=2, style="Table Grid")

    des_row = table_info.rows[0].cells
    des_row[0].text = "Client"
    des_row[1].text = 'AMF'

    var_row = table_info.rows[1].cells
    var_row[0].text = "Site"
    var_row[1].text = dico_model['site']

    var_per = table_info.rows[2].cells
    var_per[0].text = "Description"
    var_per[1].text = dico_model['description']


    for cell in table_info.columns[0].cells:
        cell.width = Inches(1)

    ### Synthèse
    p = doc.add_paragraph()
    doc.add_heading("Synthèse",level=1)

    table_synth= doc.add_table(rows=8, cols=2, style="Table Grid")

    var_row = table_synth.rows[0].cells
    var_row[0].text = "Variable modélisée"
    var_row[1].text = dico_model['tag_name'] + ' (' + dico_model['tag_unit'] + ')'


    var_freq = table_synth.rows[1].cells
    var_freq[0].text = "Fréquence"
    var_freq[1].text = dico_model['freq']


    var_per = table_synth.rows[2].cells
    var_per[0].text = "Période"
    var_per[1].text = "Du " + json_rep['dataframe_info']['start_date'] + " Au " + json_rep['dataframe_info']['end_date']


    var_nbp = table_synth.rows[3].cells
    var_nbp[0].text = "Nombre de points"
    var_nbp[1].text = str(json_rep['dataframe_info']['cleaning_info']['line_count_after'])


    var_type = table_synth.rows[4].cells
    var_type[0].text = "Type de modèle"
    var_type[1].text = json_rep['model_info']['model_type']

    var_r2 = table_synth.rows[5].cells
    var_r2[0].text = "Coefficient R2"

    if json_rep['model_info']['r2_test'] is not None:
        var_r2[1].text = "Apprentissage {v1:5.1f} (%) Validation {v2:5.1f} (%)".format(v1=100.0*json_rep['model_info']['r2_train'], v2=100.0*json_rep['model_info']['r2_test'])
    else:
        var_r2[1].text = "{v1:5.1f} (%)".format(v1=100.0*json_rep['model_info']['r2_train'])


    row = table_synth.rows[6].cells
    row[0].text = "Erreur moy. relative"

    if json_rep['model_info']['mape_test'] is not None:
        row[1].text = "Apprentissage {v1:5.1f} (%) Validation {v2:5.1f} (%)".format(v1=json_rep['model_info']['mape_train'], v2=json_rep['model_info']['mape_test'])
    else:
        row[1].text = "{v1:5.1f} (%)".format(v1=json_rep['model_info']['mape_train'])


    row = table_synth.rows[7].cells
    row[0].text = "Erreur standard"

    if json_rep['model_info']['standard_deviation_test'] is not None:
        row[1].text = "Apprentissage {v1:5.1f} Validation {v2:5.1f}".format(v1=json_rep['model_info']['standard_deviation_train'], v2=json_rep['model_info']['standard_deviation_test'])
    else:
        row[1].text = "{v1:5.1f}".format(v1=json_rep['model_info']['standard_deviation_train'])



    for cell in table_synth.columns[0].cells:
        cell.width = Inches(2)

    for cell in table_synth.columns[1].cells:
        cell.width = Inches(4)

    p = doc.add_paragraph()


    p_fig = doc.add_paragraph('Modèle en fonction de la Mesure')


    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #doc.add_heading('Modèle .vs. Mesure', 2)
    doc.add_picture(dico_infos['images']['Modèle en fonction mesure.png'], width=Mm(110))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    p = doc.add_paragraph()
    doc.add_heading("Facteurs",level=1)

    ## Facteurs

    table_fac = doc.add_table(rows=len(dico_model['facteurs'])+1, cols=4, style="Table Grid")

    heading_row = table_fac.rows[0].cells

    # add headings
    heading_row[0].text = "Nom"
    heading_row[1].text = "Description"
    heading_row[2].text = "Unité"
    heading_row[3].text = "Facteur retenu"


    num_row = 0
    for f in dico_model['facteurs'].values():
        if f['used']:
            num_row = num_row + 1
            var_fact = table_fac.rows[num_row].cells
            var_fact[0].text = f['nom']
            var_fact[1].text = f['description']
            var_fact[2].text = f['unit']
            var_fact[3].text = 'Oui'

    for f in dico_model['facteurs'].values():
        if not(f['used']):
            num_row = num_row + 1
            var_fact = table_fac.rows[num_row].cells
            var_fact[0].text = f['nom']
            var_fact[1].text = f['description']
            var_fact[2].text = f['unit']
            var_fact[3].text = 'Non' 

    make_rows_bold(table_fac.rows[0])


    doc.add_heading('Distribution facteurs numériques', 2)
    doc.add_picture(dico_infos['images']['fig_histo_num.png'], width=Mm(130))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    if 'fig_histo_cat.png' in dico_infos['images'].keys():
        doc.add_heading('Répartition facteurs discrets', 2)
        doc.add_picture(dico_infos['images']['fig_histo_cat.png'], width=Mm(130))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Coefficients de Corrélation', 2)
    doc.add_picture(dico_infos['images']['fig_correlogramme.png'], width=Mm(130))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.add_heading('Graphes des corrélations Facteurs/Variable modélisée', 2)
    doc.add_picture(dico_infos['images']['fig_correlation.png'], width=Mm(130))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if 'interp' in dico_infos.keys():

        doc.add_heading("Interprétation",level=1)

        doc.add_heading('Importance globale des facteurs', 2)
        doc.add_picture(dico_infos['interp']['importance_glob.png'], width=Mm(130))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Importance détaillée des facteurs', 2)
        doc.add_picture(dico_infos['interp']['importance_det.png'], width=Mm(130))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(dico_infos['rapport']['model_info']['formula']) > 0:
        formule = dico_infos['rapport']['dataframe_info']['target']['name'] + ' = ' + dico_infos['rapport']['model_info']['formula']
        doc.add_heading("Formule",level=1)
        doc.add_paragraph(formule)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)



    return doc